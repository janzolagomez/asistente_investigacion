import streamlit as st
import pandas as pd
import google.generativeai as genai
import time
from docx import Document # Importar la librería python-docx
from io import BytesIO # Para manejar archivos en memoria

# Configuración de la página
st.set_page_config(page_title="Asistente para Matriz de Investigación", layout="wide")

# ==============================================================================
# EXPLICACIONES DETALLADAS POR PASO Y TIPO DE INVESTIGACIÓN
# ==============================================================================
explanations = {
    'tipo_investigacion': "La investigación cualitativa busca comprender fenómenos desde la perspectiva de los participantes, mientras que la cuantitativa busca medir y probar hipótesis a través de datos numéricos y análisis estadístico. La investigación mixta combina ambos enfoques para una comprensión más completa. Elige el enfoque que mejor se adapte a tu pregunta y objetivos.",
    'tema': {
        'Cualitativa': "El tema en investigación cualitativa es una idea general que expresa el fenómeno, sujetos, actores y contexto que quieres estudiar, enfocado en la comprensión profunda. Debe ser amplio pero delimitado.",
        'Cuantitativa': "El tema en investigación cuantitativa debe ser específico, delimitado, e incluir al menos las variables principales y el contexto de estudio. Se enfoca en la medición y la relación entre variables.",
        'Mixta': "El tema en investigación mixta debe ser lo suficientemente amplio para integrar componentes cualitativos y cuantitativos, expresando el fenómeno y el contexto desde ambas perspectivas para una comprensión holística."
    },
    'pregunta': {
        'Cualitativa': "La pregunta cualitativa es una pregunta amplia y abierta que expresa el fenómeno principal que se desea comprender desde la perspectiva de los participantes, sin buscar medir o cuantificar.",
        'Cuantitativa': "La pregunta cuantitativa es una formulación clara, específica y objetiva que plantea una relación, efecto, diferencia o nivel entre una o más variables medibles. Guía la recolección y el análisis estadístico de datos.",
        'Mixta': "La pregunta de investigación mixta integra componentes cualitativos y cuantitativos, buscando comprender un fenómeno de manera más profunda, combinando la exploración de significados con la medición de relaciones o impactos."
    },
    'objetivo_general': {
        'Cualitativa': "En la investigación cualitativa, el objetivo general busca orientar la exploración, comprensión, descripción o interpretación del fenómeno o experiencia en un grupo social o comunidad específica, de manera coherente con un enfoque interpretativo. **Debe iniciar con un verbo en infinitivo (ejemplos: comprender, explorar, describir, interpretar, analizar, investigar).**",
        'Cuantitativa': "En la investigación cuantitativa, el objetivo general debe expresar claramente qué se quiere analizar, correlacionar, describir o explicar en términos de la relación, efecto o influencia entre las variables de estudio, en una población y contexto definidos. **Debe iniciar con un verbo en infinitivo (ejemplos: analizar, determinar, evaluar, establecer, comparar, medir).**",
        'Mixta': "En la investigación mixta, el objetivo general busca integrar la comprensión cualitativa y la medición cuantitativa para abordar el fenómeno de estudio de manera comprehensiva. **Debe iniciar con un verbo en infinitivo que refleje la integración (ejemplos: explorar y determinar, comprender y evaluar, analizar la relación e interpretar las percepciones).**"
    },
    'objetivos_especificos': {
        'Cualitativa': "Son metas concretas y delimitadas que el estudio busca alcanzar para lograr el objetivo general. En cualitativa, suelen enfocarse en acciones como identificar, analizar, describir o interpretar dimensiones del fenómeno.",
        'Cuantitativa': "Son metas concretas y medibles que derivan del objetivo general, detallando los pasos para alcanzarlo. En cuantitativa, se enfocan en describir variables, comparar grupos, correlacionar variables o explicar relaciones específicas.",
        'Mixta': "Son metas concretas que desglosan el objetivo general, incluyendo tanto pasos cualitativos (explorar, comprender, interpretar) como cuantitativos (medir, cuantificar, comparar, correlacionar) para abordar el fenómeno desde ambas perspectivas. Deben reflejar la secuencialidad o concurrencia de las fases."
    },
    'variables.independiente': "Es la característica o propiedad observable y medible que se presume es la *causa* o el factor que influye en otra variable. El investigador la manipula o mide para observar su efecto.",
    'variables.dependiente': "Es la característica o propiedad observable y medible que se presume es el *efecto* o el resultado que cambia debido a la influencia de la variable independiente. Es lo que se observa o mide como respuesta.",
    'hipotesis.nula': "La hipótesis nula (H₀) es una afirmación que postula la ausencia de relación, diferencia o efecto entre variables. Se asume verdadera hasta que los datos demuestren lo contrario.",
    'hipotesis.alternativa': "La hipótesis alternativa (H₁) es la afirmación que el investigador busca establecer. Contradice la hipótesis nula, sugiriendo la existencia de una relación, efecto o diferencia significativa entre las variables.",
    'justificacion': "La justificación explica la *importancia* y el *porqué* de tu investigación. Debe argumentar su relevancia teórica (qué aporta al conocimiento), práctica (cómo resuelve un problema) y social (a quién beneficia o impacta positivamente).",
    'marco_teorico': {
        'Cualitativa': "El marco teórico en investigación cualitativa es una síntesis y selección de **conceptos clave y temas relevantes** que fundamentan tu perspectiva del fenómeno. Sirve para construir tus categorías iniciales o 'lentes interpretativos' antes o durante la recolección de datos.",
        'Cuantitativa': "El marco teórico en investigación cuantitativa es la conceptualización formal de tus variables, basada en la literatura científica existente. Define qué significa cada variable desde un punto de vista académico o técnico, usando autores y modelos reconocidos, y guía la operacionalización y medición. **En esta etapa, concéntrate en los conceptos clave o temas de tu estudio.**",
        'Mixta': "El marco teórico en investigación mixta debe integrar conceptos y teorías relevantes de ambos enfoques (cualitativo y cuantitativo) para fundamentar la comprensión integral del fenómeno. Puede incluir conceptualizaciones de variables y categorías de análisis."
    },
    'metodologia.poblacion': "La población es el *conjunto total* de todas las personas, objetos o elementos que poseen una o más características comunes y que son el universo de tu estudio. Es el grupo al cual deseas generalizar tus hallazgos.",
    'metodologia.muestra': "La muestra es un *subconjunto representativo* de la población, seleccionado para realizar el estudio. Se describe el tipo de muestreo (probabilístico/no probabilístico), el tamaño de la muestra y los criterios de selección utilizados para garantizar que sea adecuada y permita inferencias si es cuantitativa.",
    'metodologia.tecnicas': {
        'Cualitativa': "Las técnicas de recolección de datos cualitativas son los procedimientos y herramientas que te permiten obtener información detallada y profunda para comprender el fenómeno. Ejemplos incluyen entrevistas, grupos focales, observación participante, o análisis documental.",
        'Cuantitativa': "Las técnicas de recolección de datos cuantitativas son los procedimientos y herramientas que te permiten obtener datos numéricos y estructurados para medir variables y probar hipótesis. Ejemplos incluyen encuestas con cuestionarios estandarizados, escalas de medición (Likert), o la recopilación de datos de registros existentes.",
        'Mixta': "Las técnicas de recolección de datos mixtas combinan procedimientos cualitativos (ej., entrevistas, grupos focales) y cuantitativos (ej., encuestas, escalas estandarizadas) para recopilar información rica y variada, buscando la complementariedad y triangulación de los datos."
    },
    'metodologia.filosofia': {
        'Cualitativa': "La filosofía de la investigación es la postura epistemológica sobre cómo se concibe el conocimiento y la realidad. Para la investigación cualitativa, los enfoques comunes son el Interpretativismo (que busca comprender el significado subjetivo de las experiencias) y el Pragmatismo (que se centra en la utilidad práctica del conocimiento).",
        'Cuantitativa': "La filosofía de la investigación es la postura epistemológica sobre cómo se concibe el conocimiento y la realidad. Para la investigación cuantitativa, los enfoques comunes son el Positivismo (que busca leyes generales y objetivas a través de la observación empírica y la verificación de hipótesis, buscando leyes generales).",
        'Mixta': "La filosofía de la investigación para un enfoque mixto es comúnmente el Pragmatismo, que valora la utilidad del conocimiento y la resolución de problemas, permitiendo la combinación de distintas perspectivas para lograr un objetivo de investigación más amplio y profundo."
    },
    'metodologia.enfoque': {
        'Cualitativa': "El enfoque de la investigación es el tipo de razonamiento que guía el proceso investigativo. En la investigación cualitativa, el enfoque es principalmente Inductivo, lo que significa que se parte de la observación de datos específicos para construir teorías o patrones generales.",
        'Cuantitativa': "El enfoque de la investigación es el tipo de razonamiento que guía el proceso investigativo. En la investigación cuantitativa, el enfoque es principalmente Deductivo, lo que significa que se parte de teorías o hipótesis generales para probarlas a través de la recolección y análisis de datos específicos.",
        'Mixta': "El enfoque de la investigación mixta combina razonamiento deductivo e inductivo, utilizando ambos enfoques en diferentes fases (secuencial) o de manera simultánea (concurrente), buscando la complementariedad en la construcción del conocimiento."
    },
    'metodologia.tipologia_estudio': {
        'Cualitativa': "La tipología o alcance de estudio clasifica la investigación cualitativa según su propósito y profundidad. Algunos tipos comunes incluyen: Fenomenológico (explora experiencias vividas), Hermenéutico (interpreta textos o símbolos), Crítico (analiza el poder y la injusticia), y Narrativo (examina historias de vida).",
        'Cuantitativa': "La tipología o alcance de estudio clasifica la investigación cuantitativa según su propósito. Los tipos comunes son: Descriptivo (describe características de una población), Correlacional (examina la relación entre variables), y Explicativo (busca causas y efectos de fenómenos).",
        'Mixta': "La tipología o alcance de estudio mixto depende del diseño específico, buscando integrar la exploración y la explicación. Incluye diseños como Exploratorio Secuencial (cual-cuant), Explicatorio Secuencial (cuant-cual), y Convergente Paralelo (cual+cuant)."
    },
    'metodologia.horizonte_tiempo': {
        'Cualitativa': "El horizonte de tiempo se refiere al plazo temporal del estudio en función de su duración y momentos de observación. Puede ser Transversal (los datos se recogen en un único momento) o Longitudinal (los datos se recogen en múltiples momentos a lo largo del tiempo).",
        'Cuantitativa': "El horizonte de tiempo se refiere al plazo temporal del estudio en función de su duración y momentos de observación. Puede ser Transversal (los datos se recogen en un único momento) o Longitudinal (los datos se recogen en múltiples momentos a lo largo del tiempo).",
        'Mixta': "El horizonte de tiempo se refiere al plazo temporal del estudio en función de su duración y momentos de observación. Puede ser Transversal (los datos cualitativos y cuantitativos se recogen en un único momento, simultáneamente) o Longitudinal (los datos cualitativos y/o cuantitativos se recogen en múltiples momentos, a lo largo del tiempo, permitiendo observar la evolución)."
    },
    'metodologia.estrategias': {
        'Cualitativa': "Las estrategias de investigación cualitativa son los diseños estructurales generales para abordar el estudio. Ejemplos incluyen: Estudio de caso, Investigación Acción Participativa (IAP), Etnográfico y Teoría Fundamentada. Cada una ofrece una forma particular de acercarse al fenómeno para una comprensión profunda.",
        'Cuantitativa': "Las estrategias de investigación cuantitativa son los diseños estructurales generales que se emplean para la recolección y análisis de datos numéricos. Ejemplos comunes son el Diseño de Encuesta, Experimental, Cuasi-experimental y No experimental. Cada estrategia define cómo se manipularán o se observarán las variables y cómo se recolectarán los datos.",
        'Mixta': "Las estrategias de investigación mixta son diseños que integran explícitamente los componentes cualitativos y cuantitativos. Incluyen diseños como Convergente Paralelo, Exploratorio Secuencial (QUAL-quan) y Explicatorio Secuencial (QUAN-qual), los cuales definen la secuencia y la forma de integración de los datos."
    }
}

# ==============================================================================
# PROMPTS PARA LA VALIDACIÓN CON GEMINI (REAL)
# ==============================================================================
gemini_prompts = {
    'tipo_investigacion': lambda respuesta: f"""
Actúa como un experto en metodología de investigación. Evalúa la elección del tipo de investigación '{respuesta}'.

Estructura tu respuesta en:
1. Reconocimiento del aporte del estudiante.
2. Evaluación crítica fundamentada: ¿el tipo de investigación es coherente con el enfoque general del estudio?
3. Orientación para la mejora (si aplica).
4. Ejemplo orientativo (si aplica).

Extensión máxima: 300 tokens. Mantén un tono académico, respetuoso y crítico.
""",
    'tema': {
        'Cualitativa': lambda tema: f"""
Actúa como experto en investigación cualitativa. Evalúa el siguiente tema de investigación:

"{tema}"

Estructura tu respuesta en:
1. Valoración inicial del esfuerzo.
2. Evaluación crítica: ¿delimita fenómeno y contexto? ¿es apropiado para estudio cualitativo?
3. Sugerencias claras de mejora.
4. Ejemplo orientador (no resolver).

Extensión máxima: 300 tokens. Sé claro y empático.
""",
        'Cuantitativa': lambda tema: f"""
Actúa como experto en investigación cuantitativa. Evalúa el siguiente tema:

"{tema}"

Estructura tu evaluación en:
1. Reconocimiento del aporte.
2. Evaluación crítica: ¿incluye variables? ¿es específico? ¿coherente con lo cuantitativo?
3. Orientación para mejorar.
4. Ejemplo ilustrativo (si aplica).

Responde en tono académico y constructivo. Extensión máxima: 300 tokens.
""",
        'Mixta': lambda tema: f"""
Actúa como experto en investigación mixta. Evalúa el siguiente tema de investigación:

"{tema}"

Estructura tu respuesta en:
1. Valoración inicial del esfuerzo.
2. Evaluación crítica: ¿delimita el fenómeno desde perspectivas cualitativas y cuantitativas? ¿Es lo suficientemente amplio para un diseño mixto?
3. Sugerencias claras de mejora.
4. Ejemplo orientador (no resolver).

Extensión máxima: 300 tokens. Sé claro y empático.
"""
    },
    'pregunta': {
        'Cualitativa': lambda pregunta: f"""
Eres experto en investigación cualitativa. Evalúa la siguiente pregunta:

"{pregunta}"

Tu retroalimentación debe:
1. Reconocer el esfuerzo.
2. Evaluar si es abierta, interpretativa y fenomenológica.
3. Orientar si requiere mejoras.
4. Incluir ejemplo similar como guía.

Sé crítico y empático. Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda pregunta: f"""
Actúa como experto en investigación cuantitativa. Evalúa:

"{pregunta}"

Tu evaluación debe:
1. Valorar el intento del estudiante.
2. Evaluar claridad, relación de variables, objetividad.
3. Orientar sin reemplazar.
4. Dar ejemplo comparativo.

Responde de forma crítica y constructiva. Extensión máxima: 300 tokens.
""",
        'Mixta': lambda pregunta: f"""
Eres experto en investigación mixta. Evalúa la siguiente pregunta:

"{pregunta}"

Tu retroalimentación debe:
1. Reconocer el esfuerzo.
2. Evaluar si integra componentes cualitativos y cuantitativos. ¿Es clara, específica y abarcadora para un diseño mixto?
3. Orientar si requiere mejoras.
4. Incluir ejemplo similar como guía.

Sé crítico y empático. Extensión máxima: 300 tokens.
"""
    },
    'objetivo_general': {
        'Cualitativa': lambda obj: f"""
Eres especialista en investigación cualitativa. Evalúa el objetivo general:

"{obj}"

Organiza tu respuesta en:
1. Reconocimiento.
2. Evaluación: ¿verbo en infinitivo adecuado? ¿coherente con lo cualitativo?
3. Recomendaciones claras.
4. Ejemplo tipo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda obj: f"""
Actúa como experto en metodología cuantitativa. Evalúa:

"{obj}"

Responde en:
1. Valoración inicial.
2. Evaluación técnica: ¿verbo de acción medible? ¿variables claras?
3. Orientación pedagógica.
4. Modelo orientador.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda obj: f"""
Eres especialista en investigación mixta. Evalúa el objetivo general:

"{obj}"

Organiza tu respuesta en:
1. Reconocimiento.
2. Evaluación: ¿El verbo en infinitivo refleja la integración cualitativa y cuantitativa? ¿Es coherente con un diseño mixto?
3. Recomendaciones claras.
4. Ejemplo tipo.

Extensión máxima: 300 tokens.
"""
    },
    'objetivos_especificos': {
        'Cualitativa': lambda objs: f"""
Evalúa los siguientes objetivos específicos cualitativos:

"{objs}"

Tu respuesta debe incluir:
1. Aprecio por el esfuerzo.
2. Evaluación crítica: ¿derivan del objetivo general? ¿son coherentes con lo cualitativo?
3. Orientación concreta.
4. Ejemplo orientativo parcial.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda objs: f"""
Evalúa los siguientes objetivos específicos cuantitativos:

"{objs}"

Organiza la retroalimentación en:
1. Reconocimiento inicial.
2. Evaluación crítica: ¿son medibles? ¿alineados con variables y objetivo general?
3. Recomendaciones formativas.
4. Ejemplo ilustrativo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda objs: f"""
Evalúa los siguientes objetivos específicos mixtos:

"{objs}"

Tu respuesta debe incluir:
1. Aprecio por el esfuerzo.
2. Evaluación crítica: ¿derivan del objetivo general? ¿Integran pasos cualitativos y cuantitativos? ¿Reflejan la secuencialidad/concurrencia del diseño?
3. Orientación concreta.
4. Ejemplo orientativo parcial.

Extensión máxima: 300 tokens.
"""
    },
    'variables.independiente': lambda var: f"""
Evalúa la siguiente variable independiente:

"{var}"

Estructura tu respuesta en:
1. Apreciación inicial.
2. Evaluación crítica: ¿es la causa? ¿está bien definida y operacionalizada?
3. Orientación pedagógica.
4. Ejemplo similar.

Extensión máxima: 300 tokens.
""",
    'variables.dependiente': lambda var: f"""
Evalúa la siguiente variable dependiente:

"{var}"

Organiza tu retroalimentación en:
1. Valoración del aporte.
2. Evaluación crítica: ¿representa el efecto? ¿es medible y coherente?
3. Recomendación para refinar.
4. Ejemplo modelo.

Extensión máxima: 300 tokens.
""",
    'hipotesis.nula': lambda hip: f"""
Evalúa la siguiente hipótesis nula:

"{hip}"

Sigue esta estructura:
1. Reconocimiento del esfuerzo.
2. Evaluación: ¿representa ausencia de relación/efecto? ¿es verificable?
3. Sugerencias.
4. Ejemplo orientador.

Extensión máxima: 300 tokens.
""",
    'hipotesis.alternativa': lambda hip: f"""
Evalúa la siguiente hipótesis alternativa:

"{hip}"

Desarrolla tu retroalimentación en:
1. Apreciación del intento.
2. Evaluación crítica: ¿contradice a la nula? ¿establece relación o efecto verificable?
3. Sugerencia de mejora.
4. Ejemplo ilustrativo.

Extensión máxima: 300 tokens.
""",
    'justificacion': lambda just: f"""
Evalúa la siguiente justificación:

"{just}"

Tu evaluación debe:
1. Reconocer aspectos positivos.
2. Evaluar: ¿aborda conveniencia, relevancia social, valor teórico, utilidad?
3. Orientación formativa.
4. Preguntas guía para revisión.

Extensión máxima: 300 tokens.
""",
    'marco_teorico': lambda temas: f"""
Evalúa la lista de conceptos para el marco teórico:

"{temas}"

1. Breve introducción en español.
2. Evaluación de pertinencia.
3. Genera lista de 5-10 palabras clave en inglés para búsqueda científica (Scopus, WoS).

Extensión máxima: 300 tokens.
""",
    'metodologia.poblacion': lambda pob: f"""
Evalúa la descripción de población:

"{pob}"

1. Valoración del esfuerzo.
2. Evaluación crítica: ¿está bien delimitada? ¿se identifican características comunes?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
""",
    'metodologia.muestra': lambda mue: f"""
Evalúa la muestra propuesta:

"{mue}"

1. Reconocimiento.
2. Evaluación: ¿tipo de muestreo y tamaño adecuados?
3. Orientación para ajustes.
4. Ejemplo similar.

Extensión máxima: 300 tokens.
""",
    'metodologia.tecnicas': {
        'Cualitativa': lambda tec: f"""
Evalúa técnicas e instrumentos:

"{tec}"

1. Aprecio inicial.
2. Evaluación crítica: ¿permiten recolectar los datos necesarios según el enfoque?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda tec: f"""
Evalúa técnicas e instrumentos:

"{tec}"

1. Aprecio inicial.
2. Evaluación crítica: ¿permiten recolectar los datos necesarios según el enfoque?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda tec: f"""
Evalúa las técnicas e instrumentos propuestos para un estudio mixto:

"{tec}"

1. Aprecio inicial.
2. Evaluación crítica: ¿Las técnicas e instrumentos cualitativos y cuantitativos son apropiados para el diseño mixto? ¿Se complementan para la triangulación de datos?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
"""
    },
    'metodologia.filosofia': {
        'Cualitativa': lambda filo: f"""
Evalúa la filosofía de investigación cualitativa:

"{filo}"

1. Reconocimiento del intento.
2. Evaluación: ¿se alinea con paradigmas interpretativos/pragmáticos?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda filo: f"""
Evalúa la filosofía de investigación cuantitativa:

"{filo}"

1. Apreciación inicial.
2. Evaluación: ¿se alinea con paradigma positivista/pragmático?
3. Orientación.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda filo: f"""
Evalúa la filosofía de investigación para un estudio mixto:

"{filo}"

1. Reconocimiento del intento.
2. Evaluación: ¿Es la filosofía adecuada para integrar ambos enfoques?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
"""
    },
    'metodologia.enfoque': {
        'Cualitativa': lambda enfoque: f"""
Evalúa el enfoque cualitativo:

"{enfoque}"

1. Reconocimiento.
2. Evaluación crítica: ¿se alinea con razonamiento inductivo?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda enfoque: f"""
Evalúa el enfoque cuantitativo:

"{enfoque}"

1. Apreciación.
2. Evaluación crítica: ¿se alinea con razonamiento deductivo?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda enfoque: f"""
Evalúa el enfoque para un estudio mixto:

"{enfoque}"

1. Reconocimiento.
2. Evaluación crítica: ¿Refleja la combinación de razonamiento inductivo y deductivo apropiada para un diseño mixto?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
"""
    },
    'metodologia.tipologia_estudio': {
        'Cualitativa': lambda tipologia: f"""
Evalúa la tipología del estudio cualitativo:

"{tipologia}"

1. Aprecio inicial.
2. Evaluación crítica: ¿es una clasificación reconocida? ¿coherente con el propósito?
3. Recomendación.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda tipologia: f"""
Evalúa la tipología del estudio cuantitativo:

"{tipologia}"

1. Reconocimiento.
2. Evaluación: ¿es adecuada para lo que se quiere medir o comparar?
3. Sugerencia.
4. Modelo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda tipologia: f"""
Evalúa la tipología del estudio mixto:

"{tipologia}"

1. Reconocimiento.
2. Evaluación: ¿Es una clasificación reconocida para diseños mixtos? ¿Es coherente con el propósito de integrar ambos enfoques?
3. Sugerencia.
4. Modelo.

Extensión máxima: 300 tokens.
"""
    },
    'metodologia.horizonte_tiempo': {
        'Cualitativa': lambda hor: f"""
Evalúa el horizonte de tiempo de tu estudio:

"{hor}"

1. Reconocimiento del intento.
2. Evaluación: ¿Es apropiado para la naturaleza de tu estudio cualitativo?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda hor: f"""
Evalúa el horizonte de tiempo de tu estudio:

"{hor}"

1. Reconocimiento del intento.
2. Evaluación: ¿Es apropiado para la naturaleza de tu estudio cuantitativo?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda hor: f"""
Evalúa el horizonte de tiempo de tu estudio mixto:

"{hor}"

1. Reconocimiento del intento.
2. Evaluación: ¿Es apropiado para la naturaleza de tu estudio mixto y la recolección de datos cualitativos y cuantitativos?
3. Sugerencias.
4. Ejemplo orientativo.

Extensión máxima: 300 tokens.
"""
    },
    'metodologia.estrategias': {
        'Cualitativa': lambda est: f"""
Evalúa la estrategia de investigación cualitativa:

"{est}"

1. Reconocimiento.
2. Evaluación crítica: ¿Es una estrategia reconocida y coherente con el enfoque cualitativo?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Cuantitativa': lambda est: f"""
Evalúa la estrategia de investigación cuantitativa:

"{est}"

1. Reconocimiento.
2. Evaluación crítica: ¿Es una estrategia reconocida y coherente con el enfoque cuantitativo?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
""",
        'Mixta': lambda est: f"""
Evalúa la estrategia de investigación mixta:

"{est}"

1. Reconocimiento.
2. Evaluación crítica: ¿Es una estrategia reconocida para diseños mixtos y coherente con la integración de enfoques?
3. Recomendaciones.
4. Ejemplo.

Extensión máxima: 300 tokens.
"""
    },
    'final_coherence_evaluation': lambda matriz, tipo: f"""
Eres asesor experto en metodología. Evalúa esta matriz de consistencia para una investigación de tipo '{tipo}':

"{matriz}"

Tu retroalimentación debe incluir:
1. Apreciación global del trabajo.
2. Evaluación crítica parte por parte (tema, objetivos, pregunta, marco, método, hipótesis o variables si aplica).
3. Sugerencias específicas para mejorar.
4. Ejemplos ilustrativos (si aplica).
5. Evaluación global de coherencia.

Extensión: 3000 tokens. Mantén el tipo de investigación claro y constante. Usa principios de metodología de investigación.
"""
}

# ==============================================================================
# FUNCIÓN PARA LLAMAR A LA API DE GEMINI
# ==============================================================================
def get_gemini_feedback(step_key, user_response, research_type):
    """
    Realiza una llamada a la API de Gemini para obtener retroalimentación.
    """
    try:
        genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
        model = genai.GenerativeModel('gemini-2.0-flash')

        prompt_template = gemini_prompts.get(step_key)
        if not prompt_template:
            return "No hay un prompt de validación configurado para esta sección."

        if step_key == 'final_coherence_evaluation':
            current_tokens_limit = 3000
            prompt_text = prompt_template(user_response, research_type)
        elif isinstance(prompt_template, dict) and research_type in prompt_template:
            specific_prompt_func = prompt_template.get(research_type)
            if not specific_prompt_func:
                return "No hay un prompt de validación para este tipo de investigación en esta sección."
            prompt_text = specific_prompt_func(user_response)
            current_tokens_limit = 300
        else:
            prompt_text = prompt_template(user_response)
            current_tokens_limit = 300

        response = model.generate_content(
            prompt_text,
            generation_config=genai.types.GenerationConfig(
                temperature=0.7,
                max_output_tokens=current_tokens_limit
            )
        )

        return response.text

    except Exception as e:
        return f"Error al conectar con la IA: {e}. Por favor, verifica tu clave de API y tu conexión."


# ==============================================================================
# INICIALIZACIÓN DEL ESTADO DE SESIÓN
# ==============================================================================
if 'step' not in st.session_state:
    st.session_state.step = 0
if 'matrix_data' not in st.session_state:
    st.session_state.matrix_data = {
        'tipo_investigacion': '',
        'tema': '',
        'pregunta': '',
        'objetivo_general': '',
        'objetivos_especificos': ['', '', ''],
        'justificacion': '',
        'marco_teorico': [],
        'metodologia': {
            'poblacion': '',
            'muestra': '',
            'tecnicas': '',
            'filosofia': '',
            'enfoque': '',
            'tipologia_estudio': '',
            'horizonte_tiempo': '',
            'estrategias': ''
        },
        'variables': {'independiente': '', 'dependiente': ''},
        'hipotesis': {'nula': '', 'alternativa': ''}
    }
if 'ai_feedback' not in st.session_state:
    st.session_state.ai_feedback = ""
if 'validating_ai' not in st.session_state:
    st.session_state.validating_ai = False
if 'ai_feedback_final' not in st.session_state:
    st.session_state.ai_feedback_final = ""

# ==============================================================================
# DEFINICIÓN DE PASOS Y SUS PREGUNTAS/EJEMPLOS
# ==============================================================================
# Helper function to check for infinitive verbs
def starts_with_infinitive(text):
    text = text.strip().lower()
    if not text:
        return False
    first_word = text.split(' ')[0]
    return first_word.endswith('ar') or first_word.endswith('er') or first_word.endswith('ir')

# Helper function to format matrix data for AI evaluation
def format_matrix_data_for_ai(data):
    formatted_str = []

    formatted_str.append(f"Tipo de Investigación: {data.get('tipo_investigacion', 'No definido')}")
    formatted_str.append(f"Tema de Investigación: {data.get('tema', 'No definido')}")
    formatted_str.append(f"Pregunta de Investigación: {data.get('pregunta', 'No definido')}")
    formatted_str.append(f"Objetivo General: {data.get('objetivo_general', 'No definido')}")

    obj_especificos = data.get('objetivos_especificos', [])
    formatted_str.append("Objetivos Específicos:")
    if obj_especificos:
        for oe in obj_especificos:
            formatted_str.append(f"- {oe}")
    else:
        formatted_str.append("- No definidos")

    if data.get('tipo_investigacion') in ['Cuantitativa', 'Mixta']: # Variables and Hipotesis apply to Mixed too
        formatted_str.append(f"Variable Independiente: {data['variables'].get('independiente', 'No definido')}")
        formatted_str.append(f"Variable Dependiente: {data['variables'].get('dependiente', 'No definido')}")
        formatted_str.append(f"Hipótesis Nula (H₀): {data['hipotesis'].get('nula', 'No definido')}")
        formatted_str.append(f"Hipótesis Alternativa (H₁): {data['hipotesis'].get('alternativa', 'No definido')}")

    formatted_str.append(f"Justificación: {data.get('justificacion', 'No definido')}")

    marco_teorico_items = data.get('marco_teorico', [])
    formatted_str.append("Marco Teórico:")
    if marco_teorico_items:
        for item in marco_teorico_items:
            formatted_str.append(f"- {item}")
    else:
        formatted_str.append("- No definido")

    metodologia = data.get('metodologia', {})
    formatted_str.append("Metodología:")
    formatted_str.append(f"- Población: {metodologia.get('poblacion', 'No definido')}")
    formatted_str.append(f"- Muestra: {metodologia.get('muestra', 'No definido')}")
    formatted_str.append(f"- Técnicas y procedimientos/Instrumento: {metodologia.get('tecnicas', 'No definido')}")
    formatted_str.append(f"- Filosofía de la investigación: {metodologia.get('filosofia', 'No definido')}")
    formatted_str.append(f"- Enfoque de la investigación: {metodologia.get('enfoque', 'No definido')}")
    formatted_str.append(f"- Tipología/Alcance de estudio: {metodologia.get('tipologia_estudio', 'No definido')}")
    formatted_str.append(f"- Horizonte de tiempo: {metodologia.get('horizonte_tiempo', 'No definido')}")
    formatted_str.append(f"- Estrategias de investigación: {metodologia.get('estrategias', 'No definido')}")

    return "\n".join(formatted_str)


# Función para generar el documento DOCX de la matriz
def generate_docx_from_matrix(data):
    document = Document()
    document.add_heading('Matriz de Consistencia de Investigación', level=1)

    # Información General
    document.add_heading('Información General', level=2)
    document.add_paragraph(f"Tipo de Investigación: {data.get('tipo_investigacion', 'No definido')}")
    document.add_paragraph(f"Tema de Investigación: {data.get('tema', 'No definido')}")
    document.add_paragraph(f"Pregunta de Investigación: {data.get('pregunta', 'No definido')}")
    document.add_paragraph(f"Objetivo General: {data.get('objetivo_general', 'No definido')}")

    document.add_heading('Objetivos Específicos', level=2)
    obj_especificos = data.get('objetivos_especificos', [])
    if obj_especificos:
        for oe in obj_especificos:
            document.add_paragraph(f"- {oe}", style='List Bullet')
    else:
        document.add_paragraph("No definidos")

    if data.get('tipo_investigacion') in ['Cuantitativa', 'Mixta']:
        document.add_heading('Variables e Hipótesis', level=2)
        document.add_paragraph(f"Variable Independiente: {data['variables'].get('independiente', 'No definido')}")
        document.add_paragraph(f"Variable Dependiente: {data['variables'].get('dependiente', 'No definido')}")
        document.add_paragraph(f"Hipótesis Nula (H₀): {data['hipotesis'].get('nula', 'No definido')}")
        document.add_paragraph(f"Hipótesis Alternativa (H₁): {data['hipotesis'].get('alternativa', 'No definido')}")

    document.add_heading('Justificación', level=2)
    document.add_paragraph(data.get('justificacion', 'No definido'))

    document.add_heading('Marco Teórico', level=2)
    marco_teorico_items = data.get('marco_teorico', [])
    if marco_teorico_items:
        for item in marco_teorico_items:
            document.add_paragraph(f"- {item}", style='List Bullet')
    else:
        document.add_paragraph("No definido")

    document.add_heading('Metodología', level=2)
    metodologia = data.get('metodologia', {})
    document.add_paragraph(f"Población: {metodologia.get('poblacion', 'No definido')}")
    document.add_paragraph(f"Muestra: {metodologia.get('muestra', 'No definido')}")
    document.add_paragraph(f"Técnicas y procedimientos/Instrumento: {metodologia.get('tecnicas', 'No definido')}")
    document.add_paragraph(f"Filosofía de la investigación: {metodologia.get('filosofia', 'No definido')}")
    document.add_paragraph(f"Enfoque de la investigación: {metodologia.get('enfoque', 'No definido')}")
    document.add_paragraph(f"Tipología/Alcance de estudio: {metodologia.get('tipologia_estudio', 'No definido')}")
    document.add_paragraph(f"Horizonte de tiempo: {metodologia.get('horizonte_tiempo', 'No definido')}")
    document.add_paragraph(f"Estrategias de investigación: {metodologia.get('estrategias', 'No definido')}")

    # Guardar en un objeto BytesIO
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Función para generar el documento DOCX de la retroalimentación de la IA
def generate_ai_feedback_docx(feedback_text):
    document = Document()
    document.add_heading('Análisis Crítico de la Matriz de Investigación por la IA', level=1)
    document.add_paragraph(feedback_text)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


base_steps = [
    {
        'name': "Tipo de Investigación",
        'question': "¡Hola! Vamos a crear tu matriz de investigación. ¿Qué tipo de investigación realizarás?",
        'examples': {},
        'input_type': 'radio',
        'options': ['Cualitativa', 'Cuantitativa', 'Mixta'], # Added 'Mixta' option
        'key': 'tipo_investigacion',
        'validation': lambda x: x != ''
    },
    {
        'name': "Tema de Investigación",
        'question': "¿Cuál es el tema de tu investigación? Describe brevemente el fenómeno y el contexto.",
        'examples': {
            'Cuantitativa': [
                "Impacto del uso de redes sociales en el rendimiento académico de estudiantes universitarios de primer año en la Facultad de Comunicación de la Universidad X durante el ciclo 2024-II.",
                "Relación entre el estrés académico y la calidad del sueño en estudiantes de medicina de una universidad española.",
                "Efecto de un programa de intervención nutricional en los niveles de glucosa en pacientes diabéticos tipo 2 en un centro de salud urbano."
            ],
            'Cualitativa': [
                "Percepciones docentes sobre la educación para el desarrollo sostenible implementadas en el cantón Portovelo, Ecuador.",
                "Experiencias de resiliencia en mujeres migrantes venezolanas en España durante el proceso de integración laboral.",
                "Significados atribuidos por jóvenes a la participación ciudadana en contextos urbanos desfavorecidos de Madrid."
            ],
            'Mixta': [
                "Evaluación del impacto de un programa de intervención educativa en la alfabetización digital y las percepciones de los docentes sobre su efectividad en escuelas rurales.",
                "Exploración de los factores que influyen en la deserción universitaria y cuantificación de su incidencia en una cohorte específica.",
                "Análisis de la eficacia de una terapia grupal en la reducción de síntomas de ansiedad y exploración de las experiencias individuales de los participantes."
            ]
        },
        'input_type': 'text_area',
        'key': 'tema',
        'validation': lambda x: len(x) > 20
    },
    {
        'name': "Pregunta de Investigación",
        'question': "¿Cuál es tu pregunta de investigación? Asegúrate de que sea clara, específica y esté alineada con tu tema.",
        'examples': {
            'Cuantitativa': [
                "¿De qué manera el uso de redes sociales influye en el rendimiento académico de los estudiantes universitarios de primer año de la Facultad de Comunicación de la Universidad X durante el ciclo 2024-II?",
                "¿Existe una correlación significativa entre el nivel de estrés académico y la calidad del sueño reportada por estudiantes de medicina?",
                "¿Cuál es el efecto de un programa de dieta baja en carbohidratos en la reducción de los niveles de glucosa en sangre en pacientes diabéticos tipo 2?"
            ],
            'Cualitativa': [
                "¿Qué percepciones tienen los docentes de educación básica respecto a las inserciones curriculares de la educación para el desarrollo sostenible en el cantón Portovelo?",
                "¿Cómo construyen la resiliencia las mujeres migrantes venezolanas al enfrentar los desafíos de la integración laboral en España?",
                "¿Cuáles son los significados que los jóvenes de barrios desfavorecidos de Madrid atribuyen a la participación ciudadana?"
            ],
            'Mixta': [
                "¿Qué impacto tiene un programa de alfabetización digital en el nivel de competencia digital de los docentes (cuant) y cómo perciben ellos su propia transformación y los desafíos del programa (cual)?",
                "¿Cuáles son los factores socioeconómicos y académicos asociados a la deserción universitaria (cuant), y cuáles son las experiencias vividas por los estudiantes que desertaron (cual)?",
                "¿En qué medida una terapia grupal reduce los síntomas de ansiedad (cuant) y cómo experimentan los participantes los cambios emocionales y relacionales durante el proceso terapéutico (cual)?"
            ]
        },
        'input_type': 'text_area',
        'key': 'pregunta',
        'validation': lambda x: len(x) > 20 and '?' in x
    },
    {
        'name': "Objetivo General",
        'question': "Ahora escribe tu objetivo general. ¿Qué meta principal quieres lograr con tu investigación?",
        'examples': {
            'Cuantitativa': [
                "Determinar la influencia del uso de redes sociales en el rendimiento académico de los estudiantes universitarios de primer año de la Facultad de Comunicación de la Universidad X durante el ciclo 2024-II.",
                "Analizar la relación entre el estrés académico y la calidad del sueño en estudiantes de medicina de una universidad española.",
                "Evaluar el efecto de un programa de intervención nutricional en los niveles de glucosa en pacientes diabéticos tipo 2 en un centro de salud urbano."
            ],
            'Cualitativa': [
                "Comprender las percepciones de los docentes de educación básica sobre las inserciones curriculares para el desarrollo sostenible en Portovelo.",
                "Explorar las experiencias de resiliencia en mujeres migrantes venezolanas durante su integración laboral en España.",
                "Interpretar los significados atribuidos por jóvenes a la participación ciudadana en contextos urbanos desfavorecidos de Madrid."
            ],
            'Mixta': [
                "Explorar las percepciones y experiencias de los docentes sobre la alfabetización digital e identificar el impacto cuantitativo de un programa de intervención en sus niveles de competencia digital.",
                "Determinar la incidencia de factores socioeconómicos y académicos en la deserción universitaria y comprender las razones subjetivas de los estudiantes que abandonaron sus estudios.",
                "Evaluar la efectividad de una terapia grupal en la reducción de síntomas de ansiedad y explorar las narrativas personales de los participantes sobre su proceso de cambio."
            ]
        },
        'input_type': 'text_area',
        'key': 'objetivo_general',
        # Modificación aquí: solo requiere más de 20 caracteres y empezar con infinitivo
        'validation': lambda x: len(x) > 20 and starts_with_infinitive(x)
    },
    {
        'name': "Objetivos Específicos",
        'question': "Escribe hasta 3 objetivos específicos. Estos deben detallar los pasos para alcanzar tu objetivo general. Inicia cada uno con un verbo en infinitivo. Ingresa uno por línea.",
        'examples': {
            'Cuantitativa': [
                "Identificar el tiempo promedio que los estudiantes de primer año dedican al uso de redes sociales diariamente.",
                "Establecer la relación entre el tiempo de uso de redes sociales y las calificaciones obtenidas por los estudiantes.",
                "Describir las percepciones de los estudiantes sobre el impacto de las redes sociales en su concentración y estudio."
            ],
            'Cualitativa': [
                "Caracterizar las inserciones curriculares en desarrollo sostenible implementadas por los docentes.",
                "Analizar las estrategias pedagógicas empleadas por los docentes para integrar el desarrollo sostenible.",
                "Explorar los desafíos que enfrentan los docentes al implementar la educación para el desarrollo sostenible."
            ],
            'Mixta': [
                "Cuantificar la mejora en las habilidades de alfabetización digital de los docentes tras el programa.",
                "Identificar las temáticas emergentes en las percepciones de los docentes sobre los beneficios del programa.",
                "Correlacionar la participación en el programa con los cambios en las actitudes hacia la tecnología educativa."
            ]
        },
        'input_type': 'text_area',
        'key': 'objetivos_especificos',
        'special': 'list_split',
        'validation': lambda x: len(x) > 0 and all(len(line.strip()) > 10 for line in x.split('\n') if line.strip())
    },
]

# Pasos adicionales para investigación Cuantitativa (se insertan si se selecciona 'Cuantitativa' o 'Mixta')
quantitative_specific_steps = [
    {
        'name': "Variable Independiente",
        'question': "Define tu variable independiente (la causa o el factor que se manipula o se presume que influye en otra variable).",
        'examples': {
            'Cuantitativa': [
                "Uso de redes sociales (medido en horas de conexión diaria).",
                "Horas de estudio semanales (medidas en autorreporte).",
                "Participación en programa de tutorías (variable categórica: sí/no)."
            ],
            'Mixta': [
                "Intensidad de la intervención educativa (variable categórica: alta, media, baja).",
                "Horas de participación en talleres de habilidades blandas (medida cuantitativa).",
                "Exposición a contenido mediático (frecuencia de exposición)."
            ],
            'Cualitativa': []
        },
        'input_type': 'text_input',
        'key': 'variables.independiente',
        'validation': lambda x: len(x) > 5
    },
    {
        'name': "Variable Dependiente",
        'question': "Define tu variable dependiente (el efecto o el resultado que se mide y se presume que es influenciado por la variable independiente).",
        'examples': {
            'Cuantitativa': [
                "Rendimiento académico (medido por el promedio de calificaciones finales).",
                "Nivel de ansiedad ante exámenes (medido con escala validada).",
                "Tasa de abandono universitario (variable dicotómica: abandono/continúa)."
            ],
            'Mixta': [
                "Nivel de competencia digital (medido por un test estandarizado).",
                "Satisfacción laboral (medida por escala Likert).",
                "Nivel de estrés percibido (medido por escala validada)."
            ],
            'Cualitativa': []
        },
        'input_type': 'text_input',
        'key': 'variables.dependiente',
        'validation': lambda x: len(x) > 5
    },
    {
        'name': "Hipótesis Nula (H₀)",
        'question': "Escribe tu hipótesis nula (H₀). Esta es una afirmación de no efecto o no relación. Se asume verdadera hasta que la evidencia demuestre lo contrario.",
        'examples': {
            'Cuantitativa': [
                "No existe influencia significativa del uso de redes sociales en el rendimiento académico de los estudiantes universitarios de primer año de la Facultad de Comunicación de la Universidad X durante el ciclo 2024-II.",
                "No hay diferencias significativas en el nivel de ansiedad ante exámenes entre estudiantes que reciben tutorías y los que no.",
                "La edad del estudiante no se correlaciona significativamente con su tasa de abandono universitario."
            ],
            'Mixta': [
                "El programa de alfabetización digital no produce un aumento significativo en la competencia digital de los docentes.",
                "No existe una relación significativa entre la participación en talleres de habilidades blandas y el nivel de satisfacción laboral."
            ],
            'Cualitativa': []
        },
        'input_type': 'text_area',
        'key': 'hipotesis.nula',
        'validation': lambda x: len(x) > 20
    },
    {
        'name': "Hipótesis Alternativa (H₁)",
        'question': "Escribe tu hipótesis alternativa (H₁). Esta es la afirmación que el investigador busca establecer, la que contradice la hipótesis nula.",
        'examples': {
            'Cuantitativa': [
                "Existe una influencia significativa del uso de redes sociales en el rendimiento académico de los estudiantes universitarios de primer año de la Facultad de Comunicación de la Universidad X durante el ciclo 2024-II.",
                "Existen diferencias significativas en el nivel de ansiedad ante exámenes entre estudiantes que reciben tutorías y los que no.",
                "La edad del estudiante se correlaciona significativamente de forma inversa con su tasa de abandono universitario."
            ],
            'Mixta': [
                "El programa de alfabetización digital produce un aumento significativo en la competencia digital de los docentes.",
                "Existe una relación positiva y significativa entre la participación en talleres de habilidades blandas y el nivel de satisfacción laboral."
            ],
            'Cualitativa': []
        },
        'input_type': 'text_area',
        'key': 'hipotesis.alternativa',
        'validation': lambda x: len(x) > 20
    },
]

# Pasos finales (comunes para ambos tipos de investigación)
final_common_steps = [
    {
        'name': "Justificación",
        'question': "¿Por qué es importante tu investigación? Explica su relevancia académica, social o práctica, y a quién beneficiará.",
        'examples': {
            'Cuantitativa': [
                "Esta investigación es relevante socialmente al abordar cómo el uso de redes sociales impacta el rendimiento académico, beneficiando a estudiantes y universidades con estrategias de estudio. Académicamente, contribuye al campo de la pedagogía digital y el bienestar estudiantil.",
                "El estudio sobre el estrés académico y calidad del sueño es vital para la salud mental estudiantil, aportando datos que pueden guiar intervenciones universitarias y enriquecer la literatura sobre factores no cognitivos del rendimiento.",
                "Evaluar la efectividad del programa nutricional ofrecerá evidencia empírica crucial para mejorar el manejo de la diabetes tipo 2, beneficiando directamente a pacientes y profesionales de la salud, y validando un modelo de intervención en el contexto local."
            ],
            'Cualitativa': [
                "Este estudio es relevante porque busca comprender las complejidades de la implementación de la educación para el desarrollo sostenible desde la voz de los docentes, lo que puede informar políticas educativas más contextualizadas y efectivas en Portovelo.",
                "Comprender la resiliencia en mujeres migrantes es fundamental para diseñar programas de apoyo psicosocial y laboral que respondan a sus necesidades reales, contribuyendo a una integración más humana y efectiva en la sociedad de acogida.",
                "La exploración de los significados de participación ciudadana en jóvenes de barrios desfavorecidos es crucial para fomentar su empoderamiento, desafiando narrativas preestablecidas y promoviendo una inclusión social más genuina."
            ],
            'Mixta': [
                "La investigación mixta ofrece una visión integral del impacto del programa de alfabetización, combinando datos duros de efectividad con las experiencias humanas de los docentes, lo que es crucial para un diseño de programas futuros más efectivos y centrados en el usuario.",
                "Este estudio es relevante porque no solo cuantifica los factores de deserción, sino que también ofrece narrativas profundas sobre las experiencias de los estudiantes, proporcionando una base sólida para intervenciones multidimensionales en la universidad."
            ]
        },
        'input_type': 'text_area',
        'key': 'justificacion',
        'validation': lambda x: len(x) > 50
    },
    {
        'name': "Marco Teórico",
        'question': "Para el marco teórico, ingresa los temas o conceptos clave que serán la base de tu estudio. Ingresa uno por línea.",
        'examples': {
            'Cuantitativa': [
                "Redes sociales",
                "Rendimiento académico",
                "Distracción digital"
            ],
            'Cualitativa': [
                "Inserciones curriculares",
                "Educación para el Desarrollo Sostenible (EDS)",
                "Percepción docente"
            ],
            'Mixta': [
                "Alfabetización digital",
                "Percepción de eficacia",
                "Integración curricular",
                "Teorías de aprendizaje mixto"
            ]
        },
        'input_type': 'text_area',
        'key': 'marco_teorico',
        'special': 'list_split',
        'validation': lambda x: len(x) > 0 and all(line.strip() != '' for line in x.split('\n') if line.strip())
    },
    {
        'name': "Población",
        'question': "Describe la población de tu estudio (¿quiénes son el grupo completo de personas o elementos con características comunes que son objeto de tu investigación?).",
        'examples': {
            'Cualitativa': [
                "La totalidad de docentes de educación básica en el cantón Portovelo, registrados en el distrito educativo durante el período 2024-2025.",
                "Un grupo focal de madres de familia de estudiantes con necesidades especiales en la escuela urbana X, durante el año escolar 2023-2024.",
                "Los pacientes mayores de 65 años diagnosticados con depresión mayor que asisten a la consulta de salud mental en el centro de salud Y, en el último semestre."
            ],
            'Cuantitativa': [
                "Todos los estudiantes de primer año de la Facultad de Comunicación de la Universidad X inscritos en el ciclo 2024-II (aproximadamente 500 estudiantes).",
                "La población estudiantil matriculada en programas de grado de la Facultad de Medicina de la Universidad Z durante el curso académico 2024-2025.",
                "Los residentes de la ciudad A mayores de 18 años, según el último censo poblacional disponible."
            ],
            'Mixta': [
                "Todos los docentes de educación primaria de 5 escuelas rurales de la región X.",
                "La población de estudiantes matriculados en los programas de grado de la Universidad Y durante el período 2023-2024."
            ]
        },
        'input_type': 'text_area',
        'key': 'metodologia.poblacion',
        'validation': lambda x: len(x) > 20
    },
    {
        'name': "Muestra",
        'question': "Describe la muestra de tu estudio (¿cómo seleccionarás a los participantes de la población y cuántos serán?).",
        'examples': {
            'Cualitativa': [
                "15 docentes de educación básica con al menos 5 años de experiencia y que hayan implementado proyectos de desarrollo sostenible, seleccionados por muestreo intencional o por conveniencia.",
                "6 madres de familia participantes en un grupo focal, seleccionadas mediante muestreo por bola de nieve a partir de contactos iniciales.",
                "8 pacientes mayores de 65 años que han completado al menos 3 meses de terapia, seleccionados por muestreo por criterio."
            ],
            'Cuantitativa': [
                "100 estudiantes seleccionados aleatoriamente de la población total (N=500), asegurando representatividad por sexo y programa de estudios mediante muestreo aleatorio simple.",
                "Una muestra estratificada de 250 estudiantes de medicina (125 por sexo) para asegurar la representatividad de la población, calculada con un nivel de confianza del 95% y un margen de error del 5%.",
                "384 ciudadanos seleccionados mediante muestreo aleatorio simple con listado telefónico, para una población infinita con un margen de error del 5% y un nivel de confianza del 95%."
            ],
            'Mixta': [
                "Una muestra probabilística de 100 docentes para la fase cuantitativa (encuesta) y una submuestra intencional de 10 docentes para entrevistas en la fase cualitativa.",
                "Para la fase cuantitativa, se utilizará un muestreo aleatorio simple de 300 estudiantes. Para la fase cualitativa, se seleccionarán 15 estudiantes que desertaron mediante muestreo por criterio."
            ]
        },
        'input_type': 'text_area',
        'key': 'metodologia.muestra',
        'validation': lambda x: len(x) > 20
    },
    {
        'name': "Técnicas y procedimientos/Instrumento",
        'question': "¿Qué técnicas e instrumentos usarás para recolectar y organizar los datos? (Ej. entrevistas, encuestas, observación).",
        'examples': {
            'Cuantitativa': [
                "Técnica: Encuesta / Instrumento: Cuestionario estandarizado (para recabar datos numéricos sobre uso de redes sociales y rendimiento percibido).",
                "Técnica: Análisis documental / Instrumento: Ficha de registro de expedientes académicos (para obtener promedios de calificaciones objetivas y tasas de abandono)."
                "Técnica: Medición psicométrica / Instrumento: Escalas de estrés o ansiedad (Escala de Estrés Percibido)."
            ],
            'Cualitativa': [
                "Técnica: Entrevistas / Instrumento: Guion de entrevistas semiestructuradas (para comprender percepciones y experiencias a profundidad).",
                "Técnica: Observación participante / Instrumento: Diario de campo, guía de observación (para documentar la dinámica de implementación de las inserciones curriculares)."
                "Técnica: Análisis de contenido / Instrumento: Matriz de análisis documental de documentos curriculares y planes de estudio (para identificar el enfoque del desarrollo sostenible)."
            ],
            'Mixta': [
                "Para la fase cuantitativa: Cuestionarios estandarizados de competencia digital. Para la fase cualitativa: Entrevistas semiestructuradas sobre experiencias con el programa.",
                "Para la fase cuantitativa: Análisis de bases de datos institucionales (notas, asistencia). Para la fase cualitativa: Grupos focales con estudiantes desertores."
            ]
        },
        'input_type': 'text_area',
        'key': 'metodologia.tecnicas',
        'validation': lambda x: len(x) > 20
    },
    # NUEVAS SECCIONES DE METODOLOGÍA
    {
        'name': "Filosofía de la investigación",
        'question': "Describe la postura epistemológica sobre cómo se concibe el conocimiento y la realidad en tu investigación.",
        'examples': {
            'Cualitativa': [
                "Interpretativismo: La realidad es una construcción social, subjetiva y múltiple, que debe ser comprendida a través de la interpretación de los significados que los individuos le dan.",
                "Pragmatismo: El conocimiento es provisional y se valida por su utilidad y las consecuencias prácticas de las acciones; se enfoca en resolver problemas."
            ],
            'Cuantitativa': [
                "Positivismo: La realidad es objetiva y externa, y el conocimiento se obtiene a través de la observación empírica y la verificación de hipótesis, buscando leyes generales.",
                "Pragmatismo: El conocimiento es provisional y se valida por su utilidad y las consecuencias prácticas de las acciones; se enfoca en resolver problemas."
            ],
            'Mixta': [
                "Pragmatismo: Se centra en la aplicación práctica del conocimiento para resolver problemas, permitiendo la combinación de enfoques y métodos de investigación para una comprensión más completa del fenómeno."
            ]
        },
        'input_type': 'text_area',
        'key': 'metodologia.filosofia',
        'validation': lambda x: len(x) > 20
    },
    {
        'name': "Enfoque de la investigación",
        'question': "Especifica el tipo de razonamiento que guía tu proceso investigativo. Selecciona la opción que mejor se adapte a tu investigación:",
        'input_type': 'radio_with_explanation',
        'options_by_type': {
            'Cualitativa': {
                "Inductivo": "Se parte de observaciones específicas y datos para desarrollar teorías, patrones y generalizaciones."
            },
            'Cuantitativa': {
                "Deductivo": "Se parte de una teoría o hipótesis general para probarla a través de observaciones específicas y datos."
            },
            'Mixta': {
                "Mixto (Secuencial o Concurrente)": "Combina razonamiento deductivo e inductivo, utilizando ambos enfoques en diferentes fases o de manera simultánea para lograr una comprensión más completa."
            }
        },
        'key': 'metodologia.enfoque',
        'validation': lambda x: x != ''
    },
    {
        'name': "Tipología/Alcance de estudio",
        'question': "Clasifica tu estudio según su propósito o alcance. Selecciona la opción que mejor se adapte a tu investigación:",
        'input_type': 'radio_with_explanation',
        'options_by_type': {
            'Cualitativa': {
                "Fenomenológico": "Busca comprender las esencias de las experiencias vividas por los individuos.",
                "Hermenéutico": "Se centra en la interpretación de textos, discursos o símbolos para comprender significados.",
                "Crítico": "Analiza las estructuras de poder y las injusticias sociales para promover el cambio.",
                "Narrativo": "Examina las historias de vida o narrativas personales para comprender fenómenos."
            },
            'Cuantitativa': {
                "Descriptivo": "Busca describir características de una población o fenómeno.",
                "Correlacional": "Examina la relación entre dos o más variables.",
                "Explicativo": "Busca establecer relaciones de causa y efecto entre variables."
            },
            'Mixta': {
                "Exploratorio Secuencial (QUAL-quan)": "Inicia con fase cualitativa para explorar y luego una cuantitativa para probar. Ejemplo: Entrevistas para generar hipótesis, luego encuesta para validarlas.",
                "Explicatorio Secuencial (QUAN-qual)": "Inicia con fase cuantitativa para obtener resultados y luego una cualitativa para profundizar o explicar. Ejemplo: Encuesta para identificar patrones, luego entrevistas para entender el 'porqué'.",
                "Concurrente (QUAL+QUAN)": "Las fases cualitativa y cuantitativa se realizan al mismo tiempo y se integran para una comprensión más completa. Ejemplo: Encuestas y grupos focales simultáneos."
            }
        },
        'key': 'metodologia.tipologia_estudio',
        'validation': lambda x: x != ''
    },
    {
        'name': "Horizonte de tiempo",
        'question': "Define el plazo temporal de tu estudio en función de su duración y momentos de observación. Selecciona la opción que mejor se adapte a tu investigación:",
        'input_type': 'radio_with_explanation',
        'options_by_type': {
            'Cualitativa': {
                "Transversal": "Los datos se recogen en un único momento, en un punto específico del tiempo.",
                "Longitudinal": "Los datos se recogen en múltiples momentos, a lo largo del tiempo, para observar cambios o desarrollo."
            },
            'Cuantitativa': {
                "Transversal": "Los datos se recogen en un único momento, en un punto específico del tiempo.",
                "Longitudinal": "Los datos se recogen en múltiples momentos, a lo largo del tiempo, para observar cambios o desarrollo."
            },
            'Mixta': {
                "Transversal": "Los datos cualitativos y cuantitativos se recogen en un único momento, simultáneamente.",
                "Longitudinal": "Los datos cualitativos y/o cuantitativos se recogen en múltiples momentos, a lo largo del tiempo, permitiendo observar la evolución."
            }
        },
        'key': 'metodologia.horizonte_tiempo',
        'validation': lambda x: x != ''
    },
    {
        'name': "Estrategias de investigación",
        'question': "Describe el diseño estructural general que emplearás para abordar tu estudio. Selecciona la opción que mejor se adapte a tu investigación:",
        'input_type': 'radio_with_explanation',
        'options_by_type': {
            'Cualitativa': {
                "Estudio de caso": "Análisis intensivo y profundo de una unidad o fenómeno específico (persona, grupo, evento).",
                "Investigación Acción Participativa (IAP)": "Proceso colaborativo de investigación y acción para resolver problemas en una comunidad.",
                "Etnográfico": "Inmersión prolongada en un entorno cultural para comprender sus prácticas y creencias.",
                "Teoría Fundamentada": "Desarrollo de una teoría a partir de los datos recopilados, sin partir de una teoría preexistente."
            },
            'Cuantitativa': {
                "Diseño de Encuesta": "Recopilación sistemática de datos de una muestra representativa para describir o analizar relaciones.",
                "Experimental": "Manipulación de una variable independiente para observar su efecto en una dependiente, con control de otras variables.",
                "Cuasi-experimental": "Similar al experimental pero sin asignación aleatoria a grupos, utilizando grupos ya existentes.",
                "No experimental": "Observación de fenómenos tal como ocurren en su contexto natural, sin manipulación de variables."
            },
            'Mixta': {
                "Diseño Convergente Paralelo": "Recoge datos cualitativos y cuantitativos de forma simultánea, los analiza por separado y luego los compara o relaciona para una comprensión integral.",
                "Diseño Exploratorio Secuencial (QUAL-quan)": "Comienza con una fase cualitativa para explorar un fenómeno, y los resultados cualitativos informan el desarrollo y la implementación de una fase cuantitativa posterior.",
                "Diseño Explicatorio Secuencial (QUAN-qual)": "Inicia con una fase cuantitativa para identificar tendencias o relaciones, y los resultados cuantitativos guían una fase cualitativa posterior para explorar las razones o explicaciones de esos resultados."
            }
        },
        'key': 'metodologia.estrategias',
        'validation': lambda x: x != ''
    },
]

# ==============================================================================
# Diccionario para nombres amigables de tipos de investigación
# ==============================================================================
tipo_invest_dict = {
    'Cualitativa': 'Cualitativa',
    'Cuantitativa': 'Cuantitativa',
    'Mixta': 'Mixta'
}

# ==============================================================================
# FUNCIÓN PRINCIPAL DE LA APLICACIÓN STREAMLIT
# ==============================================================================
def main():
    st.title("Asistente para Matriz de Investigación")
    st.write("Completa cada sección para construir tu matriz de consistencia.")
    st.markdown("---")

    # ==========================================================================
    # DETERMINACIÓN DINÁMICA DE LOS PASOS COMPLETOS
    # ==========================================================================
    tipo_investigacion = st.session_state.matrix_data.get('tipo_investigacion', '')

    all_steps = list(base_steps)
    # Variables and Hipotesis sections are added for Quantitative and Mixed types
    if tipo_investigacion in ['Cuantitativa', 'Mixta']:
        all_steps.extend(quantitative_specific_steps)
    all_steps.extend(final_common_steps)

    # ==========================================================================
    # BARRA LATERAL DE PROGRESO
    # ==========================================================================
    st.sidebar.header("Progreso de la Matriz")
    if tipo_investigacion:
        st.sidebar.markdown(f"**Tipo Seleccionado:** {tipo_invest_dict.get(tipo_investigacion, tipo_investigacion)}")
        st.sidebar.markdown("---")

    for i, step_info in enumerate(all_steps):
        icon = "⬜"
        if i < st.session_state.step:
            icon = "✅"
        elif i == st.session_state.step:
            icon = "🟨"
        st.sidebar.markdown(f"{icon} {step_info['name']}")

    # ==========================================================================
    # LÓGICA DE VISUALIZACIÓN DEL PASO ACTUAL
    # ==========================================================================
    if st.session_state.step < len(all_steps):
        current_step = all_steps[st.session_state.step]

        st.header(f"Sección: {current_step['name']}")

        # ======================================================================
        # RESUMEN DE DEFINICIONES ANTERIORES
        # ======================================================================
        friendly_names = {
            'tipo_investigacion': 'Tipo de Investigación',
            'tema': 'Tema de Investigación',
            'pregunta': 'Pregunta de Investigación',
            'objetivo_general': 'Objetivo General',
            'objetivos_especificos': 'Objetivos Específicos',
            'variables.independiente': 'Variable Independiente',
            'variables.dependiente': 'Variable Dependiente',
            'hipotesis.nula': 'Hipótesis Nula (H₀)',
            'hipotesis.alternativa': 'Hipótesis Alternativa (H₁)',
            'justificacion': 'Justificación',
            'marco_teorico': 'Marco Teórico',
            'metodologia.poblacion': 'Población',
            'metodologia.muestra': 'Muestra',
            'metodologia.tecnicas': 'Técnicas y procedimientos/Instrumento',
            'metodologia.filosofia': 'Filosofía de la investigación',
            'metodologia.enfoque': 'Enfoque de la investigación',
            'metodologia.tipologia_estudio': 'Tipología/Alcance de estudio',
            'metodologia.horizonte_tiempo': 'Horizonte de tiempo',
            'metodologia.estrategias': 'Estrategias de investigación'
        }

        completed_steps_for_summary = []
        for i in range(st.session_state.step):
            prev_step_info = all_steps[i]
            key = prev_step_info['key']

            value = None
            if '.' in key:
                main_key, sub_key = key.split('.')
                if main_key in st.session_state.matrix_data and sub_key in st.session_state.matrix_data[main_key]:
                    value = st.session_state.matrix_data[main_key][sub_key]
            else:
                value = st.session_state.matrix_data.get(key)

            if value and (isinstance(value, str) and value.strip() != '' or isinstance(value, list) and value):
                display_value = value
                if isinstance(value, list):
                    display_value = "\n".join([f"- {item}" for item in value])

                completed_steps_for_summary.append({
                    'name': friendly_names.get(key, prev_step_info['name']),
                    'value': display_value
                })

        if completed_steps_for_summary:
            with st.expander("Resumen de tus definiciones anteriores 📋"):
                for item in completed_steps_for_summary:
                    st.markdown(f"**{item['name']}:** {item['value']}")
                st.markdown("---")

        st.subheader(current_step['question'])

        exp_key = current_step['key']
        explanation_content = explanations.get(exp_key)

        if explanation_content:
            with st.expander("Ver explicación 📖"):
                if isinstance(explanation_content, dict):
                    current_research_type = st.session_state.matrix_data.get('tipo_investigacion')
                    if current_research_type:
                        st.markdown(explanation_content.get(current_research_type, "Explicación no disponible para este tipo de investigación."))
                    else:
                        st.markdown("Selecciona un tipo de investigación primero para ver la explicación relevante.")
                else:
                    st.markdown(explanation_content)

        # Removed redundant 'examples' expander for radio_with_explanation, as explanation is integrated
        if current_step.get('examples') and current_step['input_type'] not in ['radio_with_explanation', 'radio']:
            with st.expander("Ver ejemplos 💡"):
                current_research_type = st.session_state.matrix_data.get('tipo_investigacion')

                example_list = []
                if isinstance(current_step['examples'], dict):
                    if current_research_type:
                        example_list = current_step['examples'].get(current_research_type, [])
                    else:
                        st.info("Selecciona un tipo de investigación para ver los ejemplos relevantes.")
                elif isinstance(current_step['examples'], list):
                    example_list = current_step['examples']

                if example_list:
                    for i, example_text in enumerate(example_list):
                        st.markdown(f"- **Ejemplo {i+1}:** {example_text}")
                elif current_research_type and isinstance(current_step['examples'], dict):
                    st.info("No hay ejemplos específicos para este tipo de investigación en este paso.")

        st.markdown("Tu respuesta:")
        current_data_value = None
        keys = current_step['key'].split('.')
        if len(keys) == 2:
            current_data_value = st.session_state.matrix_data[keys[0]].get(keys[1], '')
        else:
            current_data_value = st.session_state.matrix_data.get(current_step['key'], '')

        # Define a function to update matrix_data and clear feedback
        def update_matrix_data_and_clear_feedback(key_to_update, new_value):
            keys_to_update_split = key_to_update.split('.')
            if len(keys_to_update_split) == 2:
                st.session_state.matrix_data[keys_to_update_split[0]][keys_to_update_split[1]] = new_value
            else:
                st.session_state.matrix_data[key_to_update] = new_value
            st.session_state.ai_feedback = "" # Clear AI feedback on data change

        if current_step['input_type'] == 'radio':
            widget_key = f"radio_input_{current_step['key']}_{st.session_state.step}"

            try:
                current_index = current_step['options'].index(current_data_value)
            except ValueError:
                current_index = 0

            response = st.radio(
                "Selecciona una opción:",
                current_step['options'],
                index=current_index,
                key=widget_key
            )

            if response != current_data_value:
                if len(keys) == 2:
                    st.session_state.matrix_data[keys[0]][keys[1]] = response
                else:
                    st.session_state.matrix_data[current_step['key']] = response
                st.session_state.ai_feedback = ""
                st.rerun()

            user_input_for_validation = response

        elif current_step['input_type'] == 'radio_with_explanation':
            current_research_type = st.session_state.matrix_data.get('tipo_investigacion')
            options_dict = {}
            if current_research_type and current_research_type in current_step['options_by_type']:
                options_dict = current_step['options_by_type'][current_research_type]

            display_options = []
            display_to_option_map = {}
            for option_name, explanation in options_dict.items():
                display_string = f"**{option_name}**: {explanation}"
                display_options.append(display_string)
                display_to_option_map[display_string] = option_name

            if display_options:
                selected_index = 0
                if current_data_value in options_dict:
                    try:
                        target_display_string = f"**{current_data_value}**: {options_dict[current_data_value]}"
                        selected_index = display_options.index(target_display_string)
                    except ValueError:
                        pass

                widget_key = f"radio_exp_input_{current_step['key']}_{st.session_state.step}"

                selected_display_option = st.radio("Selecciona una opción:", display_options,
                                                index=selected_index,
                                                key=widget_key)
                response = display_to_option_map.get(selected_display_option, "")

                if response != current_data_value:
                    if len(keys) == 2:
                        st.session_state.matrix_data[keys[0]][keys[1]] = response
                    else:
                        st.session_state.matrix_data[current_step['key']] = response
                    st.session_state.ai_feedback = ""
                    st.rerun()

                user_input_for_validation = response
            else:
                user_input_for_validation = ""
                st.warning("Selecciona primero un tipo de investigación para ver las opciones disponibles.")

        elif current_step['input_type'] == 'text_input':
            response = st.text_input("", value=current_data_value, key=f"input_{st.session_state.step}")
            if len(keys) == 2:
                st.session_state.matrix_data[keys[0]][keys[1]] = response
            else:
                st.session_state.matrix_data[current_step['key']] = response
            user_input_for_validation = response
        elif current_step['input_type'] == 'text_area':
            if current_step.get('special') == 'list_split':
                if isinstance(st.session_state.matrix_data[current_step['key']], list):
                    current_value_area = "\n".join(st.session_state.matrix_data[current_step['key']])
                else:
                    current_value_area = st.session_state.matrix_data[current_step['key']]
            else:
                current_value_area = current_data_value

            response = st.text_area("", value=current_value_area, key=f"input_{st.session_state.step}", height=150)
            user_input_for_validation = response

            if current_step.get('special') == 'list_split':
                lines = [line.strip() for line in response.split('\n') if line.strip()]
                if current_step['key'] == 'objetivos_especificos':
                    st.session_state.matrix_data[current_step['key']] = lines[:3]
                else:
                    st.session_state.matrix_data[current_step['key']] = lines
            else:
                if len(keys) == 2:
                    st.session_state.matrix_data[keys[0]][keys[1]] = response
                else:
                    st.session_state.matrix_data[current_step['key']] = response

        is_current_step_valid = current_step['validation'](user_input_for_validation)

        if not is_current_step_valid:
            if current_step['input_type'] in ['radio', 'radio_with_explanation'] and user_input_for_validation == '':
                 st.warning("Por favor, selecciona una opción para continuar.")
            elif current_step['key'] == 'tema' and len(user_input_for_validation) <= 20:
                st.warning("El tema de investigación debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'pregunta' and (len(user_input_for_validation) <= 20 or '?' not in user_input_for_validation):
                 st.warning("La pregunta debe tener al menos 20 caracteres y contener un signo de interrogación.")
            elif current_step['key'] == 'objetivo_general' and len(user_input_for_validation) <= 20:
                st.warning("El objetivo general debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'objetivo_general' and not starts_with_infinitive(user_input_for_validation):
                st.warning("El objetivo general debe empezar con un verbo en infinitivo (terminado en -ar, -er, -ir).")
            elif current_step['key'] == 'objetivos_especificos' and (len(user_input_for_validation) == 0 or not all(len(line.strip()) > 10 for line in user_input_for_validation.split('\n') if line.strip())):
                st.warning("Debes ingresar al menos un objetivo específico y cada uno debe tener al menos 10 caracteres.")
            elif current_step['key'] in ['variables.independiente', 'variables.dependiente'] and len(user_input_for_validation) <= 5:
                st.warning("El nombre de la variable debe tener al menos 5 caracteres.")
            elif current_step['key'] in ['hipotesis.nula', 'hipotesis.alternativa'] and len(user_input_for_validation) <= 20:
                st.warning("La hipótesis debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'justificacion' and len(user_input_for_validation) <= 50:
                st.warning("La justificación debe tener al menos 50 caracteres.")
            elif current_step['key'] == 'marco_teorico' and (len(user_input_for_validation) == 0 or not all(line.strip() != '' for line in user_input_for_validation.split('\n') if line.strip())):
                st.warning("Debes ingresar al menos una entrada para el marco teórico (solo los temas/conceptos).")
            elif current_step['key'] == 'metodologia.poblacion' and len(user_input_for_validation) <= 20:
                st.warning("La descripción de la población debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'metodologia.muestra' and len(user_input_for_validation) <= 20:
                st.warning("La descripción de la muestra debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'metodologia.tecnicas' and len(user_input_for_validation) <= 20:
                st.warning("La descripción de las técnicas/instrumentos debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'metodologia.filosofia' and len(user_input_for_validation) <= 20:
                st.warning("La descripción de la filosofía de investigación debe tener al menos 20 caracteres.")
            elif current_step['key'] == 'metodologia.enfoque' and user_input_for_validation == '':
                st.warning("Por favor, selecciona una opción para el enfoque de investigación.")
            elif current_step['key'] == 'metodologia.tipologia_estudio' and user_input_for_validation == '':
                st.warning("Por favor, selecciona una opción para la tipología de estudio.")
            elif current_step['key'] == 'metodologia.horizonte_tiempo' and user_input_for_validation == '':
                st.warning("Por favor, selecciona una opción para el horizonte de tiempo.")
            elif current_step['key'] == 'metodologia.estrategias' and user_input_for_validation == '':
                st.warning("Por favor, selecciona una opción para la estrategia de investigación.")
            else:
                 st.warning("Por favor, completa el campo antes de avanzar.")

        if st.button("Validar con IA ✨", disabled=not is_current_step_valid or st.session_state.validating_ai):
            st.session_state.validating_ai = True
            st.session_state.ai_feedback = ""
            with st.spinner('Validando con IA...'):
                feedback = get_gemini_feedback(
                    current_step['key'],
                    user_input_for_validation,
                    st.session_state.matrix_data.get('tipo_investigacion', '')
                )
                st.session_state.ai_feedback = feedback
            st.session_state.validating_ai = False
            st.rerun()

        if st.session_state.ai_feedback:
            st.info(st.session_state.ai_feedback)

        col1, col2 = st.columns(2)
        with col1:
            if st.session_state.step > 0:
                if st.button("⬅️ Regresar"):
                    st.session_state.step -= 1
                    st.session_state.ai_feedback = ""
                    st.rerun()
        with col2:
            if st.button("Avanzar ➡️", disabled=not is_current_step_valid):
                st.session_state.step += 1
                st.session_state.ai_feedback = ""
                st.rerun()

    else:
        st.subheader("🎉 ¡Matriz de Investigación Completa!")
        st.write("Aquí tienes un resumen de tu matriz de consistencia.")

        # Display the summary of the matrix
        data = st.session_state.matrix_data

        st.markdown("---")
        st.markdown("### Resumen de tu Matriz de Consistencia:")
        st.markdown(f"**Tipo de Investigación:** {data['tipo_investigacion'] or 'No definido'}")
        st.markdown(f"**Tema de Investigación:** {data['tema'] or 'No definido'}")
        st.markdown(f"**Pregunta de Investigación:** {data['pregunta'] or 'No definido'}")
        st.markdown(f"**Objetivo General:** {data['objetivo_general'] or 'No definido'}")

        st.markdown("**Objetivos Específicos:**")
        if data['objetivos_especificos']:
            for oe in data['objetivos_especificos']:
                if oe: st.markdown(f"- {oe}")
        else: st.markdown("No definido")

        if data['tipo_investigacion'] in ['Cuantitativa', 'Mixta']:
            st.markdown(f"**Variable Independiente:** {data['variables']['independiente'] or 'No definido'}")
            st.markdown(f"**Variable Dependiente:** {data['variables']['dependiente'] or 'No definido'}")
            st.markdown(f"**Hipótesis Nula (H₀):** {data['hipotesis']['nula'] or 'No definido'}")
            st.markdown(f"**Hipótesis Alternativa (H₁):** {data['hipotesis']['alternativa'] or 'No definido'}")

        st.markdown(f"**Justificación:** {data['justificacion'] or 'No definido'}")

        st.markdown("**Marco Teórico:**")
        if data['marco_teorico']:
            for entry in data['marco_teorico']:
                st.markdown(f"- **{entry}**")
        else: st.markdown("No definido")

        st.markdown("**Metodología:**")
        st.markdown(f"- **Población:** {data['metodologia']['poblacion'] or 'No definido'}")
        st.markdown(f"- **Muestra:** {data['metodologia']['muestra'] or 'No definido'}")
        st.markdown(f"- **Técnicas y procedimientos/Instrumento:** {data['metodologia']['tecnicas'] or 'No definido'}")
        st.markdown(f"- **Filosofía de la investigación:** {data['metodologia']['filosofia'] or 'No definido'}")
        st.markdown(f"- **Enfoque de la investigación:** {data['metodologia']['enfoque'] or 'No definido'}")
        st.markdown(f"- **Tipología/Alcance de estudio:** {data['metodologia']['tipologia_estudio'] or 'No definido'}")
        st.markdown(f"- **Horizonte de tiempo:** {data['metodologia']['horizonte_tiempo'] or 'No definido'}")
        st.markdown(f"- **Estrategias de investigación:** {data['metodologia']['estrategias'] or 'No definido'}")
        st.markdown("---")

        # Comprehensive AI Evaluation
        st.subheader("Evaluación Crítica Completa de la Matriz por la IA 🧐")
        st.write("A continuación, se evaluará la coherencia de toda tu matriz.")

        if st.button("Obtener Evaluación Crítica de la Matriz ✨"):
            st.session_state.validating_ai = True
            st.session_state.ai_feedback_final = ""
            with st.spinner('Realizando evaluación crítica de toda la matriz...'):
                formatted_matrix = format_matrix_data_for_ai(st.session_state.matrix_data)
                final_feedback = get_gemini_feedback(
                    'final_coherence_evaluation',
                    formatted_matrix,
                    st.session_state.matrix_data.get('tipo_investigacion', '')
                )
                st.session_state.ai_feedback_final = final_feedback
            st.session_state.validating_ai = False
            st.rerun()

        if st.session_state.get('ai_feedback_final'):
            st.markdown(f"**Análisis del Experto:**")
            st.info(st.session_state.ai_feedback_final)
            st.markdown("---")

            # Download button for AI feedback
            ai_feedback_doc_bytes = generate_ai_feedback_docx(st.session_state.ai_feedback_final)
            st.download_button(
                label="Descargar Análisis de la IA como DOCX �",
                data=ai_feedback_doc_bytes,
                file_name="Analisis_IA_Matriz_Investigacion.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        st.subheader("Mini Rúbrica de Autoevaluación:")
        st.write("¡Es hora de reflexionar sobre la coherencia de tu diseño!")
        st.checkbox("¿Mi pregunta de investigación está claramente alineada con mis objetivos?")
        st.checkbox("¿Mis objetivos específicos detallan los pasos para alcanzar mi objetivo general?")

        if data['tipo_investigacion'] == 'Cualitativa':
            st.checkbox("¿Las categorías analíticas son pertinentes para mi pregunta cualitativa?")
            st.checkbox("¿Las técnicas de recolección de datos son adecuadas para mi enfoque cualitativo?")
        elif data['tipo_investigacion'] == 'Cuantitativa':
            st.checkbox("¿Mis variables (independiente y dependiente) están bien definidas y son medibles?")
            st.checkbox("¿Mis hipótesis son coherentes con mis variables y tipo de estudio?")
            st.checkbox("¿La metodología que propongo es adecuada para responder mi pregunta y probar mis hipótesis?")
        elif data['tipo_investigacion'] == 'Mixta':
            st.checkbox("¿Mi diseño mixto integra coherentemente los componentes cualitativos y cuantitativos?")
            st.checkbox("¿Las preguntas y objetivos reflejan adecuadamente la naturaleza combinada del estudio?")
            st.checkbox("¿Las técnicas de recolección de datos cualitativas y cuantitativas se complementan entre sí?")

        st.markdown("---")
        st.info("¡Recuerda que este es un punto de partida! La investigación es un proceso iterativo. Lee, ajusta y perfecciona tu matriz con la literatura científica.")

        # Download button for the DOCX file of the full matrix
        if st.button("Descargar Matriz Completa como DOCX 📄"):
            docx_bytes = generate_docx_from_matrix(st.session_state.matrix_data)
            st.download_button(
                label="Haz clic aquí para descargar",
                data=docx_bytes,
                file_name="Matriz_de_Consistencia.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


        if st.button("🔄 Empezar una nueva matriz"):
            st.session_state.step = 0
            st.session_state.matrix_data = {
                'tipo_investigacion': '',
                'tema': '',
                'pregunta': '',
                'objetivo_general': '',
                'objetivos_especificos': ['', '', ''],
                'justificacion': '',
                'marco_teorico': [],
                'metodologia': {
                    'poblacion': '',
                    'muestra': '',
                    'tecnicas': '',
                    'filosofia': '',
                    'enfoque': '',
                    'tipologia_estudio': '',
                    'horizonte_tiempo': '',
                    'estrategias': ''
                },
                'variables': {'independiente': '', 'dependiente': ''},
                'hipotesis': {'nula': '', 'alternativa': ''}
            }
            st.session_state.ai_feedback = ""
            st.session_state.ai_feedback_final = ""
            st.rerun()

if __name__ == "__main__":
    main()
